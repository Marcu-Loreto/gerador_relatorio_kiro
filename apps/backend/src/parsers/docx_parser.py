"""DOCX document parser."""
import os
from typing import List

from docx import Document

from src.core.logging import get_logger
from src.graphs.state import DocumentMetadata, ExtractedSection, ExtractedTable
from src.parsers.base_parser import BaseParser, ParseResult

logger = get_logger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""
    
    def can_parse(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return file_path.lower().endswith(".docx")
    
    def parse(self, file_path: str) -> ParseResult:
        """Parse DOCX document."""
        logger.info("parsing_docx", file_path=file_path)
        
        warnings: List[str] = []
        sections: List[ExtractedSection] = []
        tables: List[ExtractedTable] = []
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            content_parts: List[str] = []
            current_section_title = "Document"
            current_section_content: List[str] = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect headings
                if para.style.name.startswith("Heading"):
                    # Save previous section
                    if current_section_content:
                        sections.append(
                            ExtractedSection(
                                title=current_section_title,
                                content="\n".join(current_section_content),
                                level=1,
                            )
                        )
                    # Start new section
                    current_section_title = text
                    current_section_content = []
                else:
                    current_section_content.append(text)
                    content_parts.append(text)
            
            # Save last section
            if current_section_content:
                sections.append(
                    ExtractedSection(
                        title=current_section_title,
                        content="\n".join(current_section_content),
                        level=1,
                    )
                )
            
            # Extract tables
            for table in doc.tables:
                rows_data: List[List[str]] = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows_data.append(row_data)
                
                if rows_data and len(rows_data) > 1:
                    tables.append(
                        ExtractedTable(
                            headers=rows_data[0],
                            rows=rows_data[1:],
                        )
                    )
            
            # Combine content
            full_content = "\n\n".join(content_parts)
            normalized_content = self.normalize_text(full_content)
            
            # Create metadata
            file_size = os.path.getsize(file_path)
            word_count = len(normalized_content.split())
            
            # Try to extract core properties
            author = None
            created_at = None
            try:
                core_props = doc.core_properties
                author = core_props.author
                created_at = core_props.created
            except Exception:
                pass
            
            metadata = DocumentMetadata(
                filename=os.path.basename(file_path),
                file_type="docx",
                file_size=file_size,
                word_count=word_count,
                author=author,
                created_at=created_at,
            )
            
            return ParseResult(
                content=normalized_content,
                metadata=metadata,
                sections=sections,
                tables=tables,
                warnings=warnings,
            )
            
        except Exception as e:
            logger.error("docx_parsing_failed", error=str(e))
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
