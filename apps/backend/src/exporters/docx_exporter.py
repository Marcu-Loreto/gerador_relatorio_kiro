"""Markdown → DOCX exporter using python-docx (no pandoc required)."""
from __future__ import annotations

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_path: str, docx_path: str | None = None) -> str:
    """Convert a markdown file to DOCX. Returns the output path."""
    if docx_path is None:
        docx_path = os.path.splitext(md_path)[0] + ".docx"
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    return md_string_to_docx(md_content, docx_path)


def md_string_to_docx(md_content: str, docx_path: str) -> str:
    """Convert markdown string to DOCX file."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    _set_default_style(doc)

    for line in md_content.splitlines():
        line = line.rstrip()

        # Headings
        if line.startswith("######"):
            _add_heading(doc, line[6:].strip(), level=6)
        elif line.startswith("#####"):
            _add_heading(doc, line[5:].strip(), level=5)
        elif line.startswith("####"):
            _add_heading(doc, line[4:].strip(), level=4)
        elif line.startswith("###"):
            _add_heading(doc, line[3:].strip(), level=3)
        elif line.startswith("##"):
            _add_heading(doc, line[2:].strip(), level=2)
        elif line.startswith("#"):
            _add_heading(doc, line[1:].strip(), level=1)

        # Horizontal rule
        elif re.match(r"^[-*_]{3,}$", line):
            _add_hr(doc)

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            _add_list_item(doc, line[2:].strip())

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            _add_list_item(doc, text.strip(), numbered=True)

        # Blockquote
        elif line.startswith("> "):
            _add_blockquote(doc, line[2:].strip())

        # Code block (single line with backticks)
        elif line.startswith("```") or line.startswith("    "):
            _add_code(doc, line.lstrip("`").lstrip())

        # Empty line → paragraph break
        elif line == "":
            doc.add_paragraph()

        # Markdown image: ![alt](data:image/png;base64,...) or ![alt](path)
        elif line.startswith("!["):
            _add_image(doc, line)

        # Normal paragraph
        else:
            _add_paragraph(doc, line)

    doc.save(docx_path)
    return docx_path


# ── Helpers ────────────────────────────────────────────────────────────────

def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def _add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 18, 2: 14, 3: 12, 4: 11, 5: 11, 6: 10}
    colors = {1: RGBColor(0x1a, 0x3a, 0x5c), 2: RGBColor(0x1a, 0x3a, 0x5c),
              3: RGBColor(0x2c, 0x52, 0x82), 4: RGBColor(0x2c, 0x52, 0x82)}
    p = doc.add_paragraph()
    run = p.add_run(_strip_inline(text))
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, RGBColor(0x1a, 0x1a, 0x1a))
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _apply_inline(p, text)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_list_item(doc: Document, text: str, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except Exception:
        p = doc.add_paragraph()
        p.add_run("• " if not numbered else "  ")
    _apply_inline(p, text)


def _add_blockquote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(_strip_inline(text))
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.left_indent = Inches(0.4)


def _add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.3)


def _add_hr(doc: Document) -> None:
    p = doc.add_paragraph("─" * 60)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)


def _add_image(doc: Document, line: str) -> None:
    """Insert image from markdown image syntax — supports base64 data URIs and file paths."""
    import re, base64, io
    m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
    if not m:
        return
    src = m.group(2).strip()
    try:
        if src.startswith("data:image"):
            # base64 embedded image
            header, b64data = src.split(",", 1)
            img_bytes = base64.b64decode(b64data)
            img_stream = io.BytesIO(img_bytes)
        elif os.path.exists(src):
            img_stream = src  # type: ignore[assignment]
        else:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(5.5))
    except Exception:
        pass  # skip broken images silently


def _strip_inline(text: str) -> str:
    """Remove markdown inline markers for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _apply_inline(p, text: str) -> None:
    """Parse inline bold/italic/code and add styled runs."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            p.add_run(text[last:m.start()])
        raw = m.group(0)
        if raw.startswith("**"):
            run = p.add_run(m.group(2))
            run.bold = True
        elif raw.startswith("*"):
            run = p.add_run(m.group(3))
            run.italic = True
        elif raw.startswith("`"):
            run = p.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        p.add_run(text[last:])
