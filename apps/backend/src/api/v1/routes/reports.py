"""Report generation, persistence and retrieval routes."""
import uuid
from typing import Optional

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse, PlainTextResponse
from pydantic import BaseModel

from src.core.logging import get_logger
from src.graphs.state import ReportType
from src.infrastructure.database import (
    count_reports,
    delete_report,
    get_report,
    list_reports,
    save_report,
)
from src.infrastructure.report_storage import (
    delete_report_file,
    read_report_file,
    save_report_file,
)

logger = get_logger(__name__)
router = APIRouter()


# ── Schemas ────────────────────────────────────────────────────────────────

class GenerateReportRequest(BaseModel):
    document_id: str
    report_type: ReportType


class ReportResponse(BaseModel):
    report_id: str
    document_id: str
    document_name: str
    report_type: str
    status: str
    quality_score: Optional[float] = None
    md_path: Optional[str] = None
    created_at: Optional[str] = None
    updated_at: Optional[str] = None
    markdown: Optional[str] = None  # populated on demand


class UpdateReportRequest(BaseModel):
    markdown: str


class ReportListResponse(BaseModel):
    total: int
    page: int
    page_size: int
    results: list[ReportResponse]


# ── Helpers ────────────────────────────────────────────────────────────────

def _to_response(row: dict, include_markdown: bool = False) -> ReportResponse:
    resp = ReportResponse(
        report_id=row["id"],
        document_id=row["document_id"],
        document_name=row["document_name"],
        report_type=row["report_type"],
        status=row["status"],
        quality_score=row.get("quality_score"),
        md_path=row.get("md_path"),
        created_at=row.get("created_at"),
        updated_at=row.get("updated_at"),
    )
    if include_markdown:
        resp.markdown = read_report_file(row.get("md_path", ""))
    return resp


# ── Generate ───────────────────────────────────────────────────────────────

@router.post("/generate", response_model=ReportResponse)
async def generate_report(request: GenerateReportRequest):
    """Generate a report from an analyzed document and persist it."""
    from src.api.v1.routes.documents import documents_store
    from src.agents.report_router import get_report_agent
    from src.graphs.state import AppState, AnalysisStatus
    from src.infrastructure.database import save_document

    if request.document_id not in documents_store:
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    doc = documents_store[request.document_id]
    if doc["status"] != "analyzed":
        raise HTTPException(status_code=400, detail="Documento precisa ser analisado primeiro")

    analysis = doc.get("analysis")
    if not analysis:
        raise HTTPException(status_code=400, detail="Dados de análise não disponíveis")

    # Persist document metadata to DB
    save_document(doc)

    report_id = str(uuid.uuid4())

    state: AppState = {
        "request_id": report_id,
        "user_id": "demo_user",
        "document_id": request.document_id,
        "document_name": doc["filename"],
        "document_type": doc["file_type"],
        "original_file_path": doc["file_path"],
        "normalized_content": analysis["content"],
        "extracted_sections": [],
        "extracted_tables": [],
        "analysis_status": AnalysisStatus.COMPLETED,
        "analysis_summary": f"Documento com {analysis['word_count']} palavras",
        "selected_report_type": request.report_type,
        "generated_report_markdown": "",
        "review_feedback": "",
        "revision_count": 0,
        "errors": [],
        "current_node": "generate_report",
    }

    try:
        from src.agents.report_router import get_report_agent
        agent = get_report_agent(request.report_type.value)
        state = agent.generate(state)
    except Exception as e:
        logger.error("report_generation_failed", error=str(e))
        raise HTTPException(status_code=500, detail=f"Falha na geração: {str(e)}")

    if state.get("errors"):
        raise HTTPException(status_code=500, detail=state["errors"][-1])

    markdown = state.get("generated_report_markdown", "")

    # Save file to disk under reports/
    md_path = save_report_file(
        report_id=report_id,
        document_name=doc["filename"],
        report_type=request.report_type.value,
        markdown=markdown,
    )

    # Persist metadata to SQLite
    from datetime import datetime
    now = datetime.utcnow().isoformat()
    record = {
        "report_id": report_id,
        "document_id": request.document_id,
        "document_name": doc["filename"],
        "report_type": request.report_type.value,
        "status": "generated",
        "quality_score": None,
        "md_path": md_path,
        "created_at": now,
    }
    save_report(record)

    logger.info("report_saved", report_id=report_id, path=md_path)

    row = get_report(report_id)
    resp = _to_response(row, include_markdown=True)
    return resp


# ── List / Search ──────────────────────────────────────────────────────────

@router.get("", response_model=ReportListResponse)
async def list_all_reports(
    search: Optional[str] = Query(None, description="Busca por nome do documento"),
    report_type: Optional[str] = Query(None, description="Filtrar por tipo de relatório"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
):
    """List and search saved reports."""
    offset = (page - 1) * page_size
    rows = list_reports(search=search, report_type=report_type, limit=page_size, offset=offset)
    total = count_reports(search=search, report_type=report_type)

    return ReportListResponse(
        total=total,
        page=page,
        page_size=page_size,
        results=[_to_response(r) for r in rows],
    )


# ── Get single ────────────────────────────────────────────────────────────

@router.get("/{report_id}", response_model=ReportResponse)
async def get_single_report(report_id: str):
    """Get report metadata + markdown content."""
    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")
    return _to_response(row, include_markdown=True)


# ── Update (user edits) ────────────────────────────────────────────────────

@router.put("/{report_id}", response_model=ReportResponse)
async def update_report(report_id: str, request: UpdateReportRequest):
    """Save user edits back to disk and DB."""
    from src.infrastructure.database import update_report_content

    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")

    # Overwrite file on disk
    md_path = row.get("md_path", "")
    if md_path:
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(request.markdown)
    else:
        # Create new file if path missing
        md_path = save_report_file(
            report_id=report_id,
            document_name=row["document_name"],
            report_type=row["report_type"],
            markdown=request.markdown,
        )

    update_report_content(report_id, md_path, status="edited")
    row = get_report(report_id)
    return _to_response(row, include_markdown=True)


# ── Delete ─────────────────────────────────────────────────────────────────

@router.delete("/{report_id}")
async def remove_report(report_id: str):
    """Delete report from DB and disk."""
    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")

    delete_report_file(row.get("md_path", ""))
    delete_report(report_id)
    return {"deleted": report_id}


# ── Export endpoints ───────────────────────────────────────────────────────

@router.get("/{report_id}/markdown")
async def get_markdown(report_id: str):
    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")
    return {"markdown": read_report_file(row.get("md_path", ""))}


@router.get("/{report_id}/export/md")
async def export_md(report_id: str):
    """Download .md file."""
    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")
    md_path = row.get("md_path", "")
    if not md_path or not __import__("os").path.exists(md_path):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado em disco")
    return FileResponse(
        md_path,
        media_type="text/markdown",
        filename=__import__("os").path.basename(md_path),
    )


@router.get("/{report_id}/export/pdf")
async def export_pdf(report_id: str):
    """Export report as PDF."""
    import asyncio
    import os
    import tempfile

    import markdown as md_lib
    from xhtml2pdf import pisa

    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")

    markdown_content = read_report_file(row.get("md_path", ""))
    if not markdown_content:
        raise HTTPException(status_code=404, detail="Conteúdo do relatório não encontrado")

    def _generate_pdf() -> str:
        html_body = md_lib.markdown(
            markdown_content,
            extensions=["tables", "fenced_code", "toc", "nl2br"],
        )
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2cm; }}
  body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11px; line-height: 1.6; color: #222; }}
  h1 {{ color: #1a365d; border-bottom: 2px solid #2b6cb0; padding-bottom: 6px; font-size: 20px; }}
  h2 {{ color: #2b6cb0; font-size: 16px; }}
  h3 {{ color: #3182ce; font-size: 14px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 10px; text-align: left; font-size: 10px; }}
  th {{ background: #edf2f7; font-weight: bold; }}
  code {{ background: #f7fafc; padding: 1px 4px; font-size: 10px; }}
  pre {{ background: #f7fafc; padding: 12px; font-size: 10px; }}
  blockquote {{ border-left: 3px solid #2b6cb0; margin: 12px 0; padding: 6px 12px; color: #555; }}
</style>
</head><body>{html_body}</body></html>"""

        os.makedirs("exports", exist_ok=True)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", dir="exports")
        pisa_status = pisa.CreatePDF(full_html, dest=tmp, encoding="utf-8")
        tmp.close()
        if pisa_status.err:
            raise RuntimeError(f"PDF generation failed with {pisa_status.err} errors")
        return tmp.name

    pdf_path = await asyncio.to_thread(_generate_pdf)
    doc_name = row.get("document_name", "relatorio")
    base_name = os.path.splitext(doc_name)[0]

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename=f"{base_name}.pdf",
    )


@router.get("/{report_id}/export/docx")
async def export_docx(report_id: str):
    """Export report as DOCX."""
    import asyncio
    import os
    import re
    import tempfile

    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")

    markdown_content = read_report_file(row.get("md_path", ""))
    if not markdown_content:
        raise HTTPException(status_code=404, detail="Conteúdo do relatório não encontrado")

    def _generate_docx() -> str:
        doc = DocxDocument()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in markdown_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            # Headings
            if stripped.startswith("### "):
                p = doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                p = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                p = doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                text = re.sub(r"^\d+\.\s", "", stripped)
                doc.add_paragraph(text, style="List Number")
            elif stripped.startswith("> "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
                run = p.add_run(stripped[2:])
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif stripped.startswith("---") or stripped.startswith("***"):
                doc.add_paragraph("─" * 50)
            else:
                # Clean bold/italic markdown
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
                text = re.sub(r"\*(.+?)\*", r"\1", text)
                text = re.sub(r"`(.+?)`", r"\1", text)
                doc.add_paragraph(text)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir="exports")
        doc.save(tmp.name)
        tmp.close()
        return tmp.name

    docx_path = await asyncio.to_thread(_generate_docx)
    doc_name = row.get("document_name", "relatorio")
    base_name = os.path.splitext(doc_name)[0]

    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{base_name}.docx",
    )
